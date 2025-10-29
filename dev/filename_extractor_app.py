"""
Simple Streamlit app to extract filenames from a folder and save to docx or md file.

Usage:
    streamlit run dev/filename_extractor_app.py --server.port 8508
"""

import streamlit as st
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt


def get_filenames_from_folder(folder_path: str) -> list[str]:
    """Extract all filenames from the given folder."""
    try:
        path = Path(folder_path)
        if not path.exists():
            return []
        
        # Get all files (not directories) from the folder
        filenames = [f.name for f in path.iterdir() if f.is_file()]
        return sorted(filenames)
    except Exception as e:
        st.error(f"Error reading folder: {e}")
        return []


def save_to_markdown(filenames: list[str], output_path: str) -> bool:
    """Save filenames to a markdown file."""
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# Extracted Filenames\n\n")
            for filename in filenames:
                f.write(f"- {filename}\n")
        return True
    except Exception as e:
        st.error(f"Error saving to markdown: {e}")
        return False


def save_to_docx(filenames: list[str], output_path: str) -> bool:
    """Save filenames to a Word docx file."""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Extracted Filenames', level=1)
        
        # Add each filename as a paragraph
        for filename in filenames:
            p = doc.add_paragraph(filename)
            # Optional: adjust font size
            run = p.runs[0]
            run.font.size = Pt(11)
        
        doc.save(output_path)
        return True
    except Exception as e:
        st.error(f"Error saving to docx: {e}")
        return False


def main():
    st.set_page_config(
        page_title="Filename Extractor",
        page_icon="📂",
        layout="wide"
    )
    
    st.title("📂 Filename Extractor")
    st.markdown("Extract filenames from a folder and save to docx or markdown file.")
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Default folder suggestion
        default_folder = "kb/StrlSchExt__db_inserted"
        
        # Folder path input
        folder_path = st.text_input(
            "Folder Path",
            value=default_folder,
            help="Enter the path to the folder containing files"
        )
        
        # Output format selection
        output_format = st.radio(
            "Output Format",
            options=["Markdown (.md)", "Word Document (.docx)"],
            index=0
        )
        
        # Output filename
        default_output = "extracted_filenames"
        output_name = st.text_input(
            "Output Filename (without extension)",
            value=default_output,
            help="Name for the output file"
        )
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📁 Source Folder")
        
        # Check if folder exists
        if folder_path:
            path = Path(folder_path)
            if path.exists() and path.is_dir():
                st.success(f"✅ Folder exists: `{folder_path}`")
                
                # Extract filenames
                filenames = get_filenames_from_folder(folder_path)
                
                if filenames:
                    st.info(f"📊 Found **{len(filenames)}** files")
                    
                    # Show preview
                    with st.expander("👀 Preview Filenames", expanded=True):
                        for idx, filename in enumerate(filenames[:50], 1):
                            st.text(f"{idx}. {filename}")
                        
                        if len(filenames) > 50:
                            st.info(f"... and {len(filenames) - 50} more files")
                else:
                    st.warning("⚠️ No files found in the folder")
            else:
                st.error(f"❌ Folder does not exist: `{folder_path}`")
                filenames = []
        else:
            st.info("ℹ️ Please enter a folder path")
            filenames = []
    
    with col2:
        st.subheader("💾 Export")
        
        if filenames:
            # Determine output path and extension
            if "Markdown" in output_format:
                output_path = f"{output_name}.md"
                file_type = "markdown"
            else:
                output_path = f"{output_name}.docx"
                file_type = "docx"
            
            st.info(f"📄 Output will be saved as: `{output_path}`")
            
            # Export button
            if st.button("💾 Extract and Save", type="primary", use_container_width=True):
                with st.spinner(f"Saving to {file_type}..."):
                    if file_type == "markdown":
                        success = save_to_markdown(filenames, output_path)
                    else:
                        success = save_to_docx(filenames, output_path)
                    
                    if success:
                        st.success(f"✅ Successfully saved {len(filenames)} filenames to `{output_path}`")
                        
                        # Provide download button
                        with open(output_path, 'rb') as f:
                            file_data = f.read()
                        
                        st.download_button(
                            label=f"⬇️ Download {output_path}",
                            data=file_data,
                            file_name=output_path,
                            mime="application/octet-stream",
                            use_container_width=True
                        )
                    else:
                        st.error("❌ Failed to save file")
        else:
            st.info("ℹ️ No files to export. Please select a valid folder with files.")
    
    # Footer
    st.markdown("---")
    st.markdown("**Usage:** Enter a folder path, select output format, and click 'Extract and Save'")


if __name__ == "__main__":
    main()
